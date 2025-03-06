import streamlit as st
import torch
import os
import time
import io
import numpy as np
import soundfile as sf
from datetime import datetime
import uuid
from docx import Document
import tempfile
import base64
import re

# Try to import accelerate - it's optional but will improve performance
try:
    import accelerate
    HAS_ACCELERATE = True
except ImportError:
    HAS_ACCELERATE = False

# Now import the transformers components
from transformers import AutoModelForSpeechSeq2Seq, AutoProcessor, pipeline

# App configuration
st.set_page_config(
    page_title="AudioTrans Pro Web",
    page_icon="🎙️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# App styling
st.markdown("""
<style>
    .main .block-container {
        padding-top: 2rem;
    }
    .stProgress .st-eb {
        background-color: #2a6099;
    }
    div[data-testid="stSidebarContent"] {
        background-color: #f8f9fa;
    }
    .transcription-result {
        max-height: 400px;
        overflow-y: auto;
        padding: 1rem;
        border: 1px solid #dee2e6;
        border-radius: 0.375rem;
        background-color: #f8f9fa;
    }
    .timestamp {
        font-weight: bold;
        color: #2a6099;
    }
    .footer {
        margin-top: 3rem;
        text-align: center;
        color: #6c757d;
    }
    .custom-info-box {
        padding: 1rem;
        border-radius: 0.375rem;
        background-color: #e3f2fd;
        border-left: 5px solid #2a6099;
        margin-bottom: 1rem;
    }
</style>
""", unsafe_allow_html=True)

# Transcriber class
class AudioTranscriber:
    def __init__(self):
        self.model = None
        self.processor = None
        self.pipe = None
        self.model_loaded = False
        
    def load_model(self):
        """Load the Whisper model and processor"""
        if self.model_loaded:
            return True
            
        try:
            # Configure device
            device = "cuda:0" if torch.cuda.is_available() else "cpu"
            torch_dtype = torch.float16 if torch.cuda.is_available() else torch.float32
            
            # Load model and processor - use a smaller model for compatibility
            model_id = "openai/whisper-small"  # Using small model
            
            # Set up model kwargs based on available optimizations - keep it minimal
            model_kwargs = {
                "torch_dtype": torch_dtype
            }
            
            # Load the model with simplified arguments
            self.model = AutoModelForSpeechSeq2Seq.from_pretrained(model_id, **model_kwargs)
            
            self.model.to(device)
            self.processor = AutoProcessor.from_pretrained(model_id)
            
            # Create the pipeline with simplified arguments
            self.pipe = pipeline(
                "automatic-speech-recognition",
                model=self.model,
                tokenizer=self.processor.tokenizer,
                feature_extractor=self.processor.feature_extractor,
                device=device
            )
            
            self.model_loaded = True
            return True
            
        except Exception as e:
            st.error(f"Error loading model: {str(e)}")
            return False
    
    def get_device_info(self):
        """Get information about the device being used"""
        if torch.cuda.is_available():
            device_name = torch.cuda.get_device_name(0)
            return f"GPU: {device_name}"
        else:
            return "Device: CPU"
    
    def get_audio_duration(self, audio_data, sampling_rate):
        """Get the duration of audio data"""
        try:
            # Calculate duration from the audio data and sample rate
            duration = len(audio_data) / sampling_rate
            return duration
        except Exception as e:
            st.warning(f"Error calculating audio duration: {str(e)}")
            return 0
    
    def transcribe(self, audio_bytes, language="en", temperature=0.0, use_timestamps=False, beam_size=1):
        """Transcribe audio from bytes"""
        if not self.model_loaded:
            success = self.load_model()
            if not success:
                return {"error": "Failed to load transcription model"}, 0
        
        try:
            start_time = time.time()
            
            # Generate kwargs for the pipeline
            generate_kwargs = {
                "task": "transcribe",
                "language": language,
            }
            
            # Only add optional parameters if they're not default values
            if temperature > 0:
                generate_kwargs["temperature"] = temperature
            if beam_size > 1:
                generate_kwargs["num_beams"] = beam_size
                
            # Check audio duration to decide if timestamps are needed
            # Calculate an approximate duration - 16000 samples = 1 second at 16kHz
            approx_duration = len(audio_bytes) / 16000  # in seconds
            
            # Enable timestamps for longer audio (>25 seconds) to avoid errors
            force_timestamps = approx_duration > 25
            actual_timestamps = use_timestamps or force_timestamps
            
            if force_timestamps and not use_timestamps:
                st.info("Note: Timestamps have been automatically enabled for this longer audio file.")
                
            # Process audio bytes directly instead of using a file path
            # This avoids the need for ffmpeg
            result = self.pipe(
                {"array": audio_bytes, "sampling_rate": 16000},  # Use 16kHz as default
                return_timestamps=actual_timestamps,
                generate_kwargs=generate_kwargs,
                chunk_length_s=30,  # Process in 30-second chunks
                stride_length_s=5   # With 5-second overlap between chunks
            )
            
            end_time = time.time()
            processing_time = end_time - start_time
            
            # Get audio duration
            audio_duration = self.get_audio_duration(audio_bytes, 16000)
            
            # Compile metadata
            metadata = {
                "date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "language": language,
                "duration": audio_duration,
                "processing_time": processing_time,
                "has_timestamps": actual_timestamps
            }
            
            # Combine result with metadata
            full_result = {
                "transcription": result,
                "metadata": metadata
            }
            
            return full_result, processing_time
            
        except Exception as e:
            return {"error": f"Transcription error: {str(e)}"}, 0
    
    def save_as_txt(self, result):
        """Save transcription as text and return the content"""
        try:
            output = io.StringIO()
            
            output.write("AUDIO TRANSCRIPTION\n")
            output.write("=" * 50 + "\n\n")
            
            # Add metadata
            metadata = result["metadata"]
            output.write(f"Transcription date: {metadata['date']}\n")
            output.write(f"Language: {metadata['language']}\n")
            output.write(f"Duration: {self.format_duration(metadata['duration'])}\n")
            output.write(f"Processing time: {self.format_duration(metadata['processing_time'])}\n\n")
            output.write("=" * 50 + "\n\n")
            
            # Add transcription
            transcription = result["transcription"]
            has_timestamps = metadata.get("has_timestamps", False)
            
            if has_timestamps and isinstance(transcription, dict) and "chunks" in transcription:
                output.write("TRANSCRIPTION WITH TIMESTAMPS:\n\n")
                for chunk in transcription["chunks"]:
                    start = chunk.get("timestamp", [0])[0]
                    minutes = int(start // 60)
                    seconds = int(start % 60)
                    time_str = f"[{minutes:02d}:{seconds:02d}] "
                    chunk_text = chunk.get("text", "")
                    output.write(time_str + chunk_text + "\n")
            else:
                if isinstance(transcription, dict) and "text" in transcription:
                    output.write(transcription["text"])
                else:
                    output.write(str(transcription))
            
            return output.getvalue()
        except Exception as e:
            st.error(f"Error creating txt file: {str(e)}")
            return None
    
    def save_as_docx(self, result):
        """Save transcription as a Word document and return the bytes"""
        try:
            # Create a new document
            doc = Document()
            
            # Add a title
            doc.add_heading('Audio Transcription', 0)
            
            # Add metadata
            metadata_table = doc.add_table(rows=4, cols=2)
            metadata_table.style = 'Table Grid'
            
            # Fill metadata
            metadata = result["metadata"]
            
            cells = metadata_table.rows[0].cells
            cells[0].text = 'Transcription date'
            cells[1].text = metadata['date']
            
            cells = metadata_table.rows[1].cells
            cells[0].text = 'Language'
            cells[1].text = metadata['language']
            
            cells = metadata_table.rows[2].cells
            cells[0].text = 'Duration'
            cells[1].text = self.format_duration(metadata['duration'])
            
            cells = metadata_table.rows[3].cells
            cells[0].text = 'Processing time'
            cells[1].text = self.format_duration(metadata['processing_time'])
            
            doc.add_paragraph('')  # Add space
            
            # Add transcription
            transcription = result["transcription"]
            has_timestamps = metadata.get("has_timestamps", False)
            
            if has_timestamps and isinstance(transcription, dict) and "chunks" in transcription:
                doc.add_heading('Transcription with timestamps', level=1)
                
                # Create table for timestamps
                timestamps_table = doc.add_table(rows=1, cols=2)
                timestamps_table.style = 'Table Grid'
                
                # Table headers
                header_cells = timestamps_table.rows[0].cells
                header_cells[0].text = 'Time'
                header_cells[1].text = 'Text'
                
                # Add chunks with timestamps
                for chunk in transcription["chunks"]:
                    start = chunk.get("timestamp", [0])[0]
                    minutes = int(start // 60)
                    seconds = int(start % 60)
                    time_str = f"{minutes:02d}:{seconds:02d}"
                    chunk_text = chunk.get("text", "")
                    
                    row_cells = timestamps_table.add_row().cells
                    row_cells[0].text = time_str
                    row_cells[1].text = chunk_text
            else:
                # Add normal transcription
                doc.add_heading('Transcription', level=1)
                
                if isinstance(transcription, dict) and "text" in transcription:
                    doc.add_paragraph(transcription["text"])
                else:
                    doc.add_paragraph(str(transcription))
            
            # Save the document to a bytes buffer
            docx_bytes = io.BytesIO()
            doc.save(docx_bytes)
            docx_bytes.seek(0)
            
            return docx_bytes.getvalue()
        except Exception as e:
            st.error(f"Error creating docx file: {str(e)}")
            return None
    
    def format_duration(self, seconds):
        """Format duration in seconds to readable format"""
        if seconds < 60:
            return f"{seconds:.1f} seconds"
        else:
            minutes = int(seconds // 60)
            secs = int(seconds % 60)
            if minutes < 60:
                return f"{minutes} minute{'s' if minutes > 1 else ''} {secs} second{'s' if secs > 1 else ''}"
            else:
                hours = int(minutes // 60)
                mins = int(minutes % 60)
                return f"{hours} hour{'s' if hours > 1 else ''} {mins} minute{'s' if mins > 1 else ''} {secs} second{'s' if secs > 1 else ''}"

# Load audio file and convert to array
def load_audio_bytes(uploaded_file):
    """Load audio bytes from uploaded file and resample to 16kHz"""
    try:
        # Read bytes from the uploaded file
        audio_bytes = uploaded_file.getvalue()
        
        # Use soundfile to read the audio data
        with sf.SoundFile(io.BytesIO(audio_bytes)) as sound_file:
            # Get the sample rate and number of channels
            sample_rate = sound_file.samplerate
            num_channels = sound_file.channels
            
            # Read the full file
            audio_data = sound_file.read(dtype='float32')
            
            # Convert to mono if needed
            if num_channels > 1:
                audio_data = audio_data.mean(axis=1)
            
            # Resample to 16kHz if needed (Whisper expects 16kHz)
            if sample_rate != 16000:
                # Simple resampling using numpy (not high quality but works)
                audio_length = len(audio_data)
                new_length = int(audio_length * 16000 / sample_rate)
                indices = np.linspace(0, audio_length - 1, new_length)
                audio_data = np.interp(indices, np.arange(audio_length), audio_data)
            
            return audio_data, 16000
    except Exception as e:
        st.error(f"Error loading audio file: {str(e)}")
        return None, None

# Create download link for file
def get_download_link(content, filename, display_text):
    """Generate a download link for the given content"""
    if filename.endswith('.txt'):
        b64 = base64.b64encode(content.encode()).decode()
        mime_type = 'text/plain'
    else:  # docx
        b64 = base64.b64encode(content).decode()
        mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        
    href = f'<a href="data:{mime_type};base64,{b64}" download="{filename}" class="download-button">{display_text}</a>'
    return href

# Initialize session state
def init_session_state():
    if 'transcriber' not in st.session_state:
        st.session_state.transcriber = AudioTranscriber()
    if 'model_loaded' not in st.session_state:
        st.session_state.model_loaded = False
    if 'transcription_result' not in st.session_state:
        st.session_state.transcription_result = None
    if 'processing_time' not in st.session_state:
        st.session_state.processing_time = 0
    if 'audio_data' not in st.session_state:
        st.session_state.audio_data = None
    if 'sample_rate' not in st.session_state:
        st.session_state.sample_rate = None

# Function to handle model loading
def load_model():
    with st.spinner('Loading Whisper model... This may take a minute or two.'):
        success = st.session_state.transcriber.load_model()
        if success:
            st.session_state.model_loaded = True
            st.success('Model loaded successfully!')
        else:
            st.error('Failed to load model. Please try again.')

# Function to handle file upload
def process_audio_file(uploaded_file):
    if uploaded_file is not None:
        with st.spinner("Processing audio file..."):
            # Load audio data
            audio_data, sample_rate = load_audio_bytes(uploaded_file)
            
            if audio_data is not None:
                # Store in session state
                st.session_state.audio_data = audio_data
                st.session_state.sample_rate = sample_rate
                
                # Calculate duration
                duration = len(audio_data) / sample_rate
                minutes = int(duration // 60)
                seconds = int(duration % 60)
                
                # Show file details
                st.success(f"File uploaded: {uploaded_file.name}")
                st.info(f"Duration: {minutes}m {seconds}s")
                
                return True
        
    return False

# Function to start transcription
def start_transcription(audio_data, language, temperature, use_timestamps, beam_size):
    if not st.session_state.model_loaded:
        st.warning("Please load the model first.")
        return
    
    if audio_data is None:
        st.warning("Please upload an audio file first.")
        return
    
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    status_text.text("Transcription in progress... Please wait.")
    
    # Simulate progress for UX
    for i in range(1, 91):
        # Slow down progress as we get closer to 90%
        if i < 60:
            time.sleep(0.01)
        else:
            time.sleep(0.03)
        progress_bar.progress(i)
    
    # Perform transcription
    try:
        result, processing_time = st.session_state.transcriber.transcribe(
            audio_data,
            language=language,
            temperature=temperature,
            use_timestamps=use_timestamps,
            beam_size=beam_size
        )
        
        st.session_state.processing_time = processing_time
        
        if 'error' in result:
            status_text.text(f"Error: {result['error']}")
            progress_bar.progress(0)
            return None
        
        # Update progress to 100%
        progress_bar.progress(100)
        status_text.text("Transcription completed successfully!")
        
        # Store the result in session state
        st.session_state.transcription_result = result
        
        return result
    
    except Exception as e:
        status_text.text(f"Error during transcription: {str(e)}")
        progress_bar.progress(0)
        return None

# Function to display transcription result
def display_transcription_result(result):
    if not result:
        return
    
    st.markdown("## Transcription Result")
    
    # Get the transcription data
    transcription = result["transcription"]
    metadata = result["metadata"]
    has_timestamps = metadata.get("has_timestamps", False)
    
    # Display processing time
    if st.session_state.processing_time > 0:
        if st.session_state.processing_time < 60:
            time_str = f"{st.session_state.processing_time:.1f} seconds"
        else:
            minutes = int(st.session_state.processing_time // 60)
            seconds = int(st.session_state.processing_time % 60)
            time_str = f"{minutes}m {seconds}s"
        
        st.info(f"Processing time: {time_str}")
    
    # Display the result
    st.markdown('<div class="transcription-result">', unsafe_allow_html=True)
    
    if has_timestamps and isinstance(transcription, dict) and "chunks" in transcription:
        # Display with timestamps
        for chunk in transcription["chunks"]:
            start = chunk.get("timestamp", [0])[0]
            minutes = int(start // 60)
            seconds = int(start % 60)
            time_str = f"[{minutes:02d}:{seconds:02d}]"
            chunk_text = chunk.get("text", "")
            
            st.markdown(f'<span class="timestamp">{time_str}</span> {chunk_text}', unsafe_allow_html=True)
    else:
        # Display normal text
        if isinstance(transcription, dict) and "text" in transcription:
            st.write(transcription["text"])
        else:
            st.write(str(transcription))
    
    st.markdown('</div>', unsafe_allow_html=True)
    
    # Create download buttons
    st.markdown("### Download Options")
    
    col1, col2 = st.columns(2)
    
    with col1:
        # Generate TXT file
        txt_content = st.session_state.transcriber.save_as_txt(result)
        if txt_content:
            filename = f"transcription_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
            st.markdown(
                get_download_link(txt_content, filename, "Download as TXT"),
                unsafe_allow_html=True
            )
    
    with col2:
        # Generate DOCX file
        docx_content = st.session_state.transcriber.save_as_docx(result)
        if docx_content:
            filename = f"transcription_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            st.markdown(
                get_download_link(docx_content, filename, "Download as DOCX"),
                unsafe_allow_html=True
            )

# Main application
def main():
    # Initialize session state
    init_session_state()
    
    # App title
    st.title("🎙️ AudioTrans Pro Web")
    st.markdown("A web-based audio transcription tool using Whisper")
    
    # Get device info
    device_info = st.session_state.transcriber.get_device_info()
    st.sidebar.info(f"💻 **Device Info**: {device_info}")
    
    # Sidebar - Model loading
    if not st.session_state.model_loaded:
        st.sidebar.warning("⚠️ The model is not loaded yet.")
        if st.sidebar.button("📥 Load Transcription Model", key="load_model", type="primary"):
            load_model()
    else:
        st.sidebar.success("✅ Model is loaded and ready to use!")
    
    # Sidebar - Settings
    st.sidebar.header("Settings")
    
    # Language selection
    language_options = {
        "English": "en",
        "French": "fr",
        "German": "de", 
        "Spanish": "es",
        "Italian": "it"
    }
    selected_language_name = st.sidebar.selectbox(
        "Language",
        options=list(language_options.keys()),
        index=0
    )
    selected_language = language_options[selected_language_name]
    
    # Output format selection (not used directly but kept for UI consistency)
    output_format = st.sidebar.selectbox(
        "Output Format",
        options=["Text (.txt)", "Word Document (.docx)"],
        index=0
    )
    
    # Precision slider
    temperature = st.sidebar.slider(
        "Precision Level",
        min_value=0.0,
        max_value=1.0,
        value=0.0,
        step=0.05,
        help="Lower values provide more precise transcription, higher values allow more creativity."
    )
    
    # Display precision label
    precision_label = "Very High"
    if temperature <= 0.2:
        precision_label = "Very High"
    elif temperature <= 0.4:
        precision_label = "High"
    elif temperature <= 0.6:
        precision_label = "Medium"
    elif temperature <= 0.8:
        precision_label = "Low"
    else:
        precision_label = "Creative"
    
    st.sidebar.caption(f"Precision: {precision_label}")
    
    # Timestamps checkbox
    use_timestamps = st.sidebar.checkbox(
        "Include timestamps",
        value=False,
        help="Add timestamps to the transcription output"
    )
    
    # Transcription quality
    quality_options = {
        "Fast": 1,
        "Standard": 2,
        "High": 3
    }
    selected_quality = st.sidebar.radio(
        "Transcription Quality",
        options=list(quality_options.keys()),
        index=0,
        help="Higher quality takes longer but may be more accurate"
    )
    beam_size = quality_options[selected_quality]
    
    # Main content - File upload
    st.header("Upload Audio File")
    
    uploaded_file = st.file_uploader(
        "Choose an audio file",
        type=["mp3", "wav", "flac", "ogg", "m4a"],
        help="Supported formats: MP3, WAV, FLAC, OGG, M4A"
    )
    
    # Process the uploaded file
    if uploaded_file is not None:
        # Check if we need to process a new file
        if st.session_state.audio_data is None or uploaded_file.name != getattr(st.session_state, 'last_uploaded_filename', None):
            st.session_state.last_uploaded_filename = uploaded_file.name
            process_audio_file(uploaded_file)
    
    # Transcribe button
    if st.button("🎬 Start Transcription", type="primary", disabled=not st.session_state.model_loaded or st.session_state.audio_data is None):
        with st.spinner("Transcribing audio..."):
            result = start_transcription(
                st.session_state.audio_data,
                selected_language,
                temperature,
                use_timestamps,
                beam_size
            )
    
    # Display transcription result if available
    if st.session_state.transcription_result:
        display_transcription_result(st.session_state.transcription_result)
    
    # Display welcome message if no transcription
    if not st.session_state.transcription_result:
        st.markdown('<div class="custom-info-box">', unsafe_allow_html=True)
        st.markdown("### 👋 Welcome to AudioTrans Pro Web!")
        st.markdown("""
        To get started:
        1. **Load the model** using the button in the sidebar
        2. **Upload an audio file** using the file uploader above
        3. **Choose settings** in the sidebar:
           - Select language
           - Adjust precision slider
           - Choose quality options
        4. **Start transcription** by clicking the button
        """)
        st.markdown("</div>", unsafe_allow_html=True)
    
    # Footer
    st.markdown(
        '<div class="footer">AudioTrans Pro Web Version - Powered by Whisper</div>',
        unsafe_allow_html=True
    )

# Run the app
if __name__ == "__main__":
    main()